import docx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- 輔助函式：設定中文字型與顏色 ---
def set_font(run, font_name='PMingLiU', size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

# --- 關鍵修正：徹底清空頁首/頁尾內容 ---
def clear_content(header_footer):
    """
    清除 Header 或 Footer 內的所有段落。
    解決 Word 預設會複製上一節內容導致內容堆疊 (11111...) 的問題。
    """
    if header_footer is None: return
    # 刪除所有現有段落
    for paragraph in header_footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    # 重新加入一個空白段落以免報錯，並回傳該段落
    return header_footer.add_paragraph()

# --- 輔助函式：加入頁碼 ---
def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    
    # 頁碼功能變數結構
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    set_font(run, size=10)

# --- 輔助函式：加入 STYLEREF (自動抓取章名) ---
def add_styleref_field(paragraph):
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # 這裡抓取 "Heading 1"
    instrText.text = 'STYLEREF "Heading 1"'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    set_font(run, size=10)

def create_system_spec_doc():
    doc = Document()
    
    # 1. 全域樣式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體')
    style.font.size = Pt(12)
    
    # 2. 封面頁 (Section 0)
    section0 = doc.sections[0]
    section0.orientation = WD_ORIENT.PORTRAIT
    section0.page_width = Cm(21.0)
    section0.page_height = Cm(29.7)
    section0.left_margin = Cm(3)
    section0.right_margin = Cm(3)
    section0.top_margin = Cm(3)
    section0.bottom_margin = Cm(3)
    
    # 封面內容
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("系統建置與規格設計說明書\n(符合直向格式規範)")
    set_font(run, size=24, bold=True, color=RGBColor(0, 0, 0))
    doc.add_page_break() # 封面結束
    
    # 啟用奇偶頁不同
    doc.settings.odd_and_even_pages_header_footer = True

    # --- 章節資料 ---
    chapters = {
        "第一章　文件概述": ["文件目的與適用範圍", "名詞定義與縮寫說明", "參考文件與相關規範"],
        "第二章　系統整體說明": ["系統建置目標", "系統使用對象與角色", "系統整體架構概述", "系統運作流程總覽"],
        "第三章　系統架構設計": ["系統分層架構（前端／後端／資料層）", "模組劃分與相依關係", "技術選型與開發環境", "系統部署架構"],
        "第四章　功能模組規格": ["AI 命題與編輯模組", "題庫管理與版本控管模組", "AI 智慧組卷模組", "測驗分析與等化模組", "帳號、權限與角色管理模組"],
        "第五章　資料庫與資料結構設計": ["關聯式資料庫設計", "知識圖譜資料結構", "資料表與欄位說明", "資料一致性與完整性規範"],
        "第六章　系統流程與介面說明": ["主要業務流程（Flow Diagram）", "使用者操作流程", "主要畫面與介面說明"],
        "第七章　資訊安全與稽核機制": ["身分驗證與存取控管", "操作紀錄與稽核追蹤", "資料安全與封閉環境設計"],
        "第八章　系統測試與驗證": ["功能測試規劃", "系統整合測試", "使用者驗收測試（UAT）"],
        "第九章　部署、維運與擴充說明": ["系統部署方式", "系統維運與備份機制", "未來擴充與升級方向"]
    }
    
    # --- 迴圈產生章節 ---
    for title, subtitles in chapters.items():
        # 新增一節 (從奇數頁開始)
        new_section = doc.add_section(WD_SECTION.ODD_PAGE)
        
        # 確保頁面設定正確 (直向)
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = Cm(21.0)
        new_section.page_height = Cm(29.7)
        new_section.left_margin = Cm(3)
        new_section.right_margin = Cm(3)
        new_section.top_margin = Cm(3)
        new_section.bottom_margin = Cm(3)
        
        # 【重要步驟】 1. 先斷開連結
        new_section.header.is_linked_to_previous = False
        new_section.even_page_header.is_linked_to_previous = False
        new_section.footer.is_linked_to_previous = False
        new_section.even_page_footer.is_linked_to_previous = False

        # 【重要步驟】 2. 徹底清空內容 (Wipe) 並取得乾淨的段落
        # 這是解決 11111 和錯誤堆疊的關鍵
        p_h_odd = clear_content(new_section.header)
        p_h_even = clear_content(new_section.even_page_header)
        p_f_odd = clear_content(new_section.footer)
        p_f_even = clear_content(new_section.even_page_footer)

        # --- 設定頁首 (Header) ---
        # 奇數頁：右側章名 (STYLEREF)
        p_h_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_styleref_field(p_h_odd)
        
        # 偶數頁：左側計畫名稱
        p_h_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_h = p_h_even.add_run("研究主題：AI 智慧測驗系統建置計畫")
        set_font(run_h, size=10)

        # --- 設定頁尾 (Footer) ---
        # 奇數頁：右側頁碼
        p_f_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(p_f_odd)

        # 偶數頁：左側頁碼
        p_f_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_page_number(p_f_even)

        # --- 寫入內容 ---
        p_title = doc.add_paragraph(title)
        p_title.style = 'Heading 1'
        set_font(p_title.runs[0], size=18, bold=True, color=RGBColor(0, 0, 0))
        
        doc.add_paragraph("【章節摘要】本章旨在說明" + title[3:] + "之核心內容與規劃重點...")
        
        for sub in subtitles:
            p_sub = doc.add_paragraph(sub)
            p_sub.style = 'Heading 2'
            set_font(p_sub.runs[0], size=14, bold=True, color=RGBColor(0, 0, 0))
            doc.add_paragraph("（內容...）\n")

    # --- 附錄 ---
    sect_app = doc.add_section(WD_SECTION.ODD_PAGE)
    sect_app.header.is_linked_to_previous = False
    sect_app.footer.is_linked_to_previous = False
    
    # 同樣要清空
    p_h_app = clear_content(sect_app.header)
    p_f_app = clear_content(sect_app.footer)
    
    p_h_app.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_styleref_field(p_h_app)
    
    p_f_app.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(p_f_app)

    p = doc.add_paragraph("附錄")
    p.style = 'Heading 1'
    set_font(p.runs[0], size=18, bold=True, color=RGBColor(0, 0, 0))
    
    appendices = ["A. 系統流程圖", "B. 資料表定義", "C. API 規格"]
    for app in appendices:
        p = doc.add_paragraph(app)
        set_font(p.runs[0], size=14, bold=True, color=RGBColor(0, 0, 0))
        doc.add_paragraph("（附錄資料）\n")

    filename = '系統規格書_Final_Fix.docx'
    doc.save(filename)
    print(f"檔案已建立：{filename}")

if __name__ == "__main__":
    create_system_spec_doc()