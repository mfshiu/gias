import docx
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn

def set_font(run, font_name='PMingLiU', size=12, bold=False, italic=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name) # 設定中文字型
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def create_report_template():
    doc = Document()
    
    # --- 全域樣式設定 ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體') # 內文指定新細明體
    style.font.size = Pt(12)
    
    # --- 頁面設定 (A4 橫向, 邊界 3cm) ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.6)  # 實際上 A4 寬度約 29.7，此處依規定設 29.6
    section.page_height = Cm(20.9) # 實際上 A4 高度約 21.0，此處依規定設 20.9
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    
    # 啟用奇偶頁不同 (為了頁首頁尾)
    section.different_first_page_header_footer = False # 封面通常獨立，這裡先設全域
    doc.settings.odd_and_even_pages_header_footer = True

    # --- 1. 封面 (Cover) ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n計畫名稱：[請輸入計畫名稱]\n")
    set_font(run, size=20, bold=True)
    
    run = p.add_run("案號：[請輸入案號]\n\n")
    set_font(run, size=16)

    run = p.add_run("國家環境研究院委託研究\n(或 國家環境研究院編印)\n\n")
    set_font(run, size=18, bold=True)
    
    run = p.add_run("計畫執行期間：民國   年   月   日至   年   月   日\n")
    set_font(run, size=14)
    run = p.add_run("受託單位：[請輸入單位名稱]\n")
    set_font(run, size=14)
    run = p.add_run("印製年月：中華民國   年   月\n")
    set_font(run, size=14)
    
    doc.add_page_break()

    # --- 1.1 封面裡頁 (Inner Cover) ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("（封面裡頁）\n\n計畫經費：新台幣           元整\n")
    set_font(run, size=14)
    run = p.add_run("受託單位計畫執行人員：[請列出人員]\n")
    set_font(run, size=14)
    doc.add_page_break()

    # --- 2. 報告基本資料表 (附件二) ---
    doc.add_paragraph("【附件二：報告基本資料表】(不編頁碼)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 3. 計畫成果中英文摘要 (簡要版) (附件三) ---
    p = doc.add_paragraph("【附件三：計畫成果中英文摘要 (簡要版)】")
    set_font(p.runs[0], bold=True)
    doc.add_paragraph("（約300-500字，不編頁碼）")
    doc.add_page_break()

    # --- 4. 目次、圖次、表次 ---
    doc.add_paragraph("目次 (TOC)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("圖次 (List of Figures)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("表次 (List of Tables)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 5. 報告大綱 ---
    doc.add_paragraph("報告大綱").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("（簡述報告架構及各章節內涵）")
    doc.add_page_break()

    # --- 6. 計畫成果摘要 (詳細版) (附件四) ---
    p = doc.add_paragraph("【附件四：計畫成果摘要 (詳細版)】")
    set_font(p.runs[0], bold=True)
    doc.add_paragraph("（約4千至1萬字，以羅馬數字編列頁碼）")
    doc.add_section(WD_SECTION.NEW_PAGE) # 新章節開始正式編碼

    # --- 7. 報告本文 (Body) ---
    # 設定頁首頁尾 (模擬)
    # 注意：python-docx 對奇偶頁首的完全控制較複雜，建議生成後在 Word 中確認
    
    chapters = ["第一章 研究緣起及目的", "第二章 執行方法", "第三章 本年度主要工作內容及成果", "第四章 結論與建議"]
    
    for chapter in chapters:
        # 確保章節自單頁(奇數頁)開始
        # Word 若要強制奇數頁分節符號，需設 section.start_type
        
        p = doc.add_paragraph(chapter)
        p.style = 'Heading 1'
        set_font(p.runs[0], size=16, bold=True)
        
        doc.add_paragraph("【章節摘要】(請撰寫於章標題後或章末)")
        
        doc.add_paragraph("內文開始...（每頁約30行，12點新細明體）")
        doc.add_paragraph("註：頁碼請註記於每頁文字下方，奇數頁頁眉為「章名」（右），偶數頁頁眉為「研究主題」（左）。")
        
        doc.add_page_break()

    # --- 8. 參考書目 ---
    p = doc.add_paragraph("參考書目")
    set_font(p.runs[0], size=16, bold=True)
    doc.add_paragraph("中文書名以粗體打印：\n王大明 (2025)。**環境工程概論**。台北：環境出版社。")
    doc.add_paragraph("西文書名以斜體打印：\nSmith, J. (2025). *Environmental Science*. New York: Academic Press.")
    doc.add_page_break()

    # --- 9. 附錄 ---
    doc.add_paragraph("附錄一：計畫成果效益自評表")
    doc.add_paragraph("附錄二：評選會議審查意見")
    doc.add_paragraph("（附錄請自單頁打印）")
    doc.add_page_break()

    # --- 10. 封底 ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("統一編號：[請填寫]").font.size = Pt(10)
    
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("＊本報告係受託單位或計畫主持人個人之意見，僅供本院施政之參考，不代表本院立場。\n")
    run = p.add_run("＊本報告之著作財產權屬國家環境研究院所有，非經本院同意，任何人均不得重製、仿製或其他之侵害。\n")
    if True: # 如有機密
        run = p.add_run("（如有屬機密或限閱者，請一併註明）")

    # Save
    file_name = "國家環境研究院_研究報告樣本.docx"
    doc.save(file_name)
    print(f"檔案已生成：{file_name}")

if __name__ == "__main__":
    create_report_template()